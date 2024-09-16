import streamlit as st
from streamlit_lottie import st_lottie
import requests
from anthropic import Anthropic
import yaml
import markdown2
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from streamlit.components.v1 import html
import base64
from bs4 import BeautifulSoup
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import re

# Set up the Anthropic client
client = Anthropic(api_key=st.secrets["ANTHROPIC_API_KEY"])

# Force light theme and set page config
st.set_page_config(page_title="EMB-AI BRD Generator", layout="wide", initial_sidebar_state="collapsed")

st.header("EMB-AI BRD Studio")

# Register Poppins fonts
pdfmetrics.registerFont(TTFont('Poppins', 'assets/Poppins-Regular.ttf'))
pdfmetrics.registerFont(TTFont('Poppins-Bold', 'assets/Poppins-Bold.ttf'))

# Function to load and encode the local image
def get_base64_of_bin_file(bin_file):
    with open(bin_file, 'rb') as f:
        data = f.read()
    return base64.b64encode(data).decode()

# Load the local watermark image
watermark_base64 = get_base64_of_bin_file('watermark.png')

# Load Lottie animations
def load_lottieurl(url: str):
    try:
        r = requests.get(url)
        if r.status_code != 200:
            return None
        return r.json()
    except:
        return None

lottie_writing = load_lottieurl("https://assets5.lottiefiles.com/packages/lf20_41X1Rp.json")
lottie_completed = load_lottieurl("https://assets4.lottiefiles.com/packages/lf20_j3yeurta.json")

def display_lottie_or_text(lottie_data, fallback_text, height=200):
    if lottie_data is not None:
        st_lottie(lottie_data, height=height)
    else:
        st.info(fallback_text)

# Load confidentiality agreement
def load_confidentiality_agreement():
    try:
        with open('confidentiality_agreement.yaml', 'r') as file:
            agreement = yaml.safe_load(file)
        
        if isinstance(agreement, dict):
            if 'Confidentiality Agreement' in agreement:
                agreement = agreement['Confidentiality Agreement']
            if 'content' in agreement:
                return agreement['content']
            elif 'title' in agreement and 'content' in agreement:
                return f"# {agreement['title']}\n\n{agreement['content']}"
        
        if isinstance(agreement, str):
            return agreement
        
        return "Confidentiality agreement structure is not recognized. Please check the YAML file."
    except FileNotFoundError:
        return "Confidentiality agreement file not found. Please ensure 'confidentiality_agreement.yaml' is in the same directory as the app."
    except yaml.YAMLError as e:
        return f"Error parsing confidentiality agreement YAML: {str(e)}"
    except Exception as e:
        return f"An unexpected error occurred while loading the confidentiality agreement: {str(e)}"

confidentiality_agreement = load_confidentiality_agreement()

# Function to count tokens
def count_tokens(text):
    try:
        encoding = tiktoken.get_encoding("cl100k_base")
        return len(encoding.encode(text))
    except Exception as e:
        print(f"Error in token counting: {str(e)}")
        return len(text.split())

# Function to determine mode based on model selections
def determine_mode(model_part1, model_part2, model_part3, model_part4):
    if all(model == "Intelligent" for model in [model_part1, model_part2, model_part3, model_part4]):
        return "Lite Mode"
    elif model_part1 == "Intelligent" and model_part2 == "Complex" and model_part3 == "Intelligent" and model_part4 == "Complex":
        return "Default Mode"
    elif all(model == "Complex" for model in [model_part1, model_part2, model_part3, model_part4]):
        return "Advanced Mode"
    else:
        return "Custom Mode"

# Function to update sidebar model selections based on mode
def update_sidebar_models(mode):
    if mode == "Lite Mode":
        st.session_state.model_part1 = "Intelligent"
        st.session_state.model_part2 = "Intelligent"
        st.session_state.model_part3 = "Intelligent"
        st.session_state.model_part4 = "Intelligent"
    elif mode == "Default Mode":
        st.session_state.model_part1 = "Intelligent"
        st.session_state.model_part2 = "Complex"
        st.session_state.model_part3 = "Intelligent"
        st.session_state.model_part4 = "Complex"
    elif mode == "Advanced Mode":
        st.session_state.model_part1 = "Complex"
        st.session_state.model_part2 = "Complex"
        st.session_state.model_part3 = "Complex"
        st.session_state.model_part4 = "Complex"

# Initialize session state for mode and models if not exists
if 'mode' not in st.session_state:
    st.session_state.mode = "Default Mode"
if 'model_part1' not in st.session_state:
    st.session_state.model_part1 = "Intelligent"
if 'model_part2' not in st.session_state:
    st.session_state.model_part2 = "Complex"
if 'model_part3' not in st.session_state:
    st.session_state.model_part3 = "Intelligent"
if 'model_part4' not in st.session_state:
    st.session_state.model_part4 = "Complex"

# Mode selection in the main area
mode = st.radio(
    "Select BRD Generation Mode:",
    ("Lite Mode", "Default Mode", "Advanced Mode", "Custom Mode"),
    index=1,  # Set Default Mode as the default selection
    key="mode_selection",
    horizontal=True 
)

# Update sidebar models if mode is changed
if mode != st.session_state.mode and mode != "Custom Mode":
    update_sidebar_models(mode)
    st.session_state.mode = mode

# Sidebar content
with st.sidebar:
    st.markdown('<div class="sidebar-content">', unsafe_allow_html=True)
    st.header("User Onboarding")
    st.info("Welcome! Please provide your information below.")
    
    user_name = st.text_input("Your Name", placeholder="John Doe")
    user_designation = st.selectbox(
        "Your Designation",
        ["Sales", "Business Analyst", "Project Manager"],
        index=0
    )
    
    if user_name or user_designation:
        welcome_message = f"Welcome, {user_name}"
        if user_designation:
            welcome_message += f" - {user_designation}"
        st.success(welcome_message)
    
    st.markdown('</div>', unsafe_allow_html=True)
    
    # Advanced Settings
    st.header("Advanced Settings")
    
    # Model selection for each part
    st.subheader("AI Model Selection")
    model_part1 = st.selectbox("Part 1 Model", ["Intelligent", "Complex"], index=0, key="model_part1")
    model_part2 = st.selectbox("Part 2 Model", ["Intelligent", "Complex"], index=1, key="model_part2")
    model_part3 = st.selectbox("Part 3 Model", ["Intelligent", "Complex"], index=0, key="model_part3")
    model_part4 = st.selectbox("Part 4 Model", ["Intelligent", "Complex"], index=1, key="model_part4")
    
    # Temperature settings for each part
    st.subheader("Temperature Settings")
    temp_part1 = st.slider("Part 1 Temperature", 0.0, 1.0, 1.0, 0.1)
    temp_part2 = st.slider("Part 2 Temperature", 0.0, 1.0, 1.0, 0.1)
    temp_part3 = st.slider("Part 3 Temperature", 0.0, 1.0, 1.0, 0.1)
    temp_part4 = st.slider("Part 4 Temperature", 0.0, 1.0, 1.0, 0.1)
    
    # Section selection for each part
    st.subheader("Section Selection")
    sections_part1 = st.multiselect(
        "Part 1 Sections",
        ["Confidentiality Agreement", "Executive Summary", "Project Approach"],
        ["Confidentiality Agreement", "Executive Summary", "Project Approach"]
    )
    sections_part2 = st.multiselect(
        "Part 2 Sections",
        ["Functional Requirements", "3rd Party Integrations and API Suggestions"],
        ["Functional Requirements", "3rd Party Integrations and API Suggestions"]
    )
    sections_part3 = st.multiselect(
        "Part 3 Sections",
        ["User Journey and User Stories", "Non-Functional Requirements"],
        ["User Journey and User Stories", "Non-Functional Requirements"]
    )
    sections_part4 = st.multiselect(
        "Part 4 Sections",
        ["Annexure"],
        ["Annexure"]
    )

    # Determine mode based on model selections
    determined_mode = determine_mode(model_part1, model_part2, model_part3, model_part4)

    # Update mode if sidebar selections change
    if determined_mode != st.session_state.mode:
        st.session_state.mode = determined_mode
        st.session_state.mode_selection = determined_mode

# Sync mode with sidebar settings
if mode != st.session_state.mode:
    st.warning(f"Mode selection doesn't match sidebar settings. Adjusting to {st.session_state.mode}.")
    mode = st.session_state.mode
    st.session_state.mode_selection = st.session_state.mode

# Function to determine model based on mode and part
def get_model(part, mode):
    if mode == "Lite Mode":
        return "claude-3-5-sonnet-20240620"
    elif mode == "Default Mode":
        if part in [1, 3]:
            return "claude-3-5-sonnet-20240620"
        else:
            return "claude-3-opus-20240229"
    elif mode == "Advanced Mode":
        return "claude-3-opus-20240229"
    else:  # Custom Mode
        model_selection = {
            1: st.session_state.model_part1,
            2: st.session_state.model_part2,
            3: st.session_state.model_part3,
            4: st.session_state.model_part4
        }
        return "claude-3-opus-20240229" if model_selection[part] == "Complex" else "claude-3-5-sonnet-20240620"

# Main content
st.markdown("""
This app generates a comprehensive Business Requirements Document (BRD) based on your inputs.
Fill in the fields below and click 'Generate BRD' to create your document.
""")

# Initialize session state for form fields
if 'form_fields' not in st.session_state:
    st.session_state.form_fields = {
        'client_name': '',
        'project_description': '',
        'user_types': '',
        'deliverables': ''
    }

# Function to update session state
def update_form_field():
    for field in st.session_state.form_fields.keys():
        if field in st.session_state:
            st.session_state.form_fields[field] = st.session_state[field]

# Progress bar
total_fields = len(st.session_state.form_fields)
filled_fields = sum(1 for value in st.session_state.form_fields.values() if value)
progress = filled_fields / total_fields
st.progress(progress)
st.write(f"Form Completion: {filled_fields}/{total_fields} fields")

# Input fields
with st.expander("Client and Project Information", expanded=True):
    client_name = st.text_input("Client Name",
                                placeholder="Enter the client's name (e.g., Acme Corporation)",
                                value=st.session_state.form_fields['client_name'],
                                key='client_name', on_change=update_form_field)
    project_description = st.text_area("Project Description and Requirements",
                                       placeholder="Provide a detailed description of the project and list the main requirements. You can use bullet points or numbered lists.",
                                       height=200,
                                       value=st.session_state.form_fields['project_description'],
                                       key='project_description', on_change=update_form_field)

with st.expander("Project Details", expanded=True):
    col1, col2 = st.columns(2)
    with col1:
        user_types = st.text_area("Types of Users",
                                  placeholder="List the different types of users who will interact with the system, one per line (e.g., Admin, Customer, Guest)",
                                  height=150,
                                  value=st.session_state.form_fields['user_types'],
                                  key='user_types', on_change=update_form_field)
    with col2:
        deliverables = st.text_area("Project Deliverables",
                                    placeholder="List the main deliverables or components of the project, one per line (e.g., User Dashboard, Admin Panel, API Integration)",
                                    height=150,
                                    value=st.session_state.form_fields['deliverables'],
                                    key='deliverables', on_change=update_form_field)

# Function to generate BRD part
def generate_brd_part(prompt, placeholder, model, temperature):
    response = ""
    
    with client.messages.stream(
        model=model,
        max_tokens=4096,
        temperature=temperature,
        messages=[
            {"role": "user", "content": prompt}
        ]
    ) as stream:
        for text in stream.text_stream:
            response += text
            placeholder.markdown(response)
    
    return response

# Convert Markdown to PDF function
def convert_markdown_to_pdf(markdown_content):
    html_content = markdown2.markdown(markdown_content, extras=["tables", "fenced-code-blocks"])

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, 
                            topMargin=2.5 * cm, bottomMargin=1.5 * cm, 
                            leftMargin=2 * cm, rightMargin=2 * cm)

    # Define fixed custom styles with correct settings
    custom_styles = {
        'CustomTitle': ParagraphStyle(
            name='CustomTitle',
            fontName='Poppins-Bold',
            fontSize=18,
            leading=22,
            alignment=TA_CENTER,
            spaceAfter=12
        ),
        'CustomHeading1': ParagraphStyle(
            name='CustomHeading1',
            fontName='Poppins-Bold',
            fontSize=16,
            leading=20,
            spaceAfter=10
        ),
        'CustomHeading2': ParagraphStyle(
            name='CustomHeading2',
            fontName='Poppins-Bold',
            fontSize=14,
            leading=18,
            spaceBefore=12,
            spaceAfter=6
        ),
        'CustomHeading3': ParagraphStyle(
            name='CustomHeading3',
            fontName='Poppins-Bold',
            fontSize=12,
            leading=16,
            spaceBefore=10,
            spaceAfter=4
        ),
        'CustomBodyText': ParagraphStyle(
            name='CustomBodyText',
            fontName='Poppins',
            fontSize=10,
            leading=14,
            alignment=TA_JUSTIFY,
            spaceAfter=8
        ),
        'CustomCode': ParagraphStyle(
            name='CustomCode',
            fontName='Courier',
            fontSize=9,
            leading=12,
            spaceAfter=8,
            leftIndent=1 * cm
        )
    }

    flowables = []

    # Function to convert HTML table to ReportLab Table
    def html_table_to_reportlab(html_table):
        soup = BeautifulSoup(html_table, 'html.parser')
        data = []
        for row in soup.find_all('tr'):
            row_data = []
            for cell in row.find_all(['td', 'th']):
                cell_content = cell.get_text(separator='\n', strip=True)
                row_data.append(Paragraph(cell_content, custom_styles['CustomBodyText']))
            data.append(row_data)

        if not data:
            return None

        table = Table(data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#11A64A')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), 'Poppins-Bold'),
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
            ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
            ('ALIGN', (0, 1), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 1), (-1, -1), 'Poppins'),
            ('FONTSIZE', (0, 1), (-1, -1), 9),
            ('TOPPADDING', (0, 1), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        return table

    # Use BeautifulSoup to parse the HTML content
    soup = BeautifulSoup(html_content, 'html.parser')

    for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'p', 'pre', 'table', 'ul', 'ol']):
        try:
            if element.name == 'h1':
                flowables.append(Paragraph(element.text, custom_styles['CustomTitle']))
            elif element.name == 'h2':
                flowables.append(Paragraph(element.text, custom_styles['CustomHeading1']))
            elif element.name == 'h3':
                flowables.append(Paragraph(element.text, custom_styles['CustomHeading2']))
            elif element.name == 'h4':
                flowables.append(Paragraph(element.text, custom_styles['CustomHeading3']))
            elif element.name == 'p':
                flowables.append(Paragraph(element.text, custom_styles['CustomBodyText']))
            elif element.name == 'pre':
                code_lines = element.text.split('\n')
                for line in code_lines:
                    flowables.append(Paragraph(line, custom_styles['CustomCode']))
            elif element.name == 'table':
                table = html_table_to_reportlab(str(element))
                if table:
                    flowables.append(table)
            elif element.name in ['ul', 'ol']:
                for li in element.find_all('li'):
                    bullet = '•' if element.name == 'ul' else f"{li.find_previous_siblings('li').__len__() + 1}."
                    text = f"{bullet} {li.text}"
                    flowables.append(Paragraph(text, custom_styles['CustomBodyText']))

            flowables.append(Spacer(1, 6))

        except Exception as e:
            print(f"Error processing element: {element}. Error: {str(e)}")
            flowables.append(Paragraph(str(element), custom_styles['CustomBodyText']))

    # Function to add watermark and page number
    def add_watermark_and_page_number(canvas, doc):
        canvas.saveState()

        # Add page number
        canvas.setFont('Poppins', 9)
        page_num = canvas.getPageNumber()
        text = f"Page {page_num}"
        canvas.drawRightString(doc.width + doc.rightMargin, doc.bottomMargin, text)

        # Add watermark to top right corner
        watermark = ImageReader('watermark.png')
        canvas.drawImage(watermark, doc.width + doc.rightMargin - 2.5*cm, doc.height + doc.topMargin - 2.5*cm, 
                         width=2*cm, height=2*cm, mask='auto', preserveAspectRatio=True)

        canvas.restoreState()

    # Build the PDF
    doc.build(flowables, onFirstPage=add_watermark_and_page_number, onLaterPages=add_watermark_and_page_number)

    buffer.seek(0)
    return buffer

# Convert Markdown to DOCX function
def convert_markdown_to_docx(markdown_content):
    doc = Document()
    
    # Set up styles
    styles = doc.styles

    # Modify the Normal style
    style_normal = styles['Normal']
    style_normal.font.name = 'Poppins'
    style_normal.font.size = Pt(10)
    style_normal.paragraph_format.space_after = Pt(8)

    # Create custom styles
    def create_style(name, font_name, font_size, bold=False, italic=False, color=RGBColor(0, 0, 0)):
        style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.name = font_name
        font.size = Pt(font_size)
        font.bold = bold
        font.italic = italic
        font.color.rgb = color
        return style

    style_title = create_style('CustomTitle', 'Poppins', 18, bold=True)
    style_heading1 = create_style('CustomHeading1', 'Poppins', 16, bold=True)
    style_heading2 = create_style('CustomHeading2', 'Poppins', 14, bold=True)
    style_heading3 = create_style('CustomHeading3', 'Poppins', 12, bold=True)
    style_code = create_style('CustomCode', 'Courier', 9)

    # Set up page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Add watermark to top right corner
    section = doc.sections[0]
    header = section.header
    watermark_paragraph = header.paragraphs[0]
    watermark_run = watermark_paragraph.add_run()
    watermark_run.add_picture("watermark.png", width=Cm(2))
    watermark_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Function to add page number
    def add_page_number(paragraph):
        run = paragraph.add_run()
        run.font.name = 'Poppins'
        run.font.size = Pt(9)
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run._element.append(fldChar1)
        run._element.append(instrText)
        run._element.append(fldChar2)

    # Add page number to footer
    section = doc.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(paragraph)

    # Parse markdown and add content to document
    html_content = markdown2.markdown(markdown_content, extras=["tables", "fenced-code-blocks"])
    soup = BeautifulSoup(html_content, 'html.parser')

    for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'p', 'pre', 'table', 'ul', 'ol']):
        try:
            if element.name == 'h1':
                doc.add_paragraph(element.text, style=style_title)
            elif element.name == 'h2':
                doc.add_paragraph(element.text, style=style_heading1)
            elif element.name == 'h3':
                doc.add_paragraph(element.text, style=style_heading2)
            elif element.name == 'h4':
                doc.add_paragraph(element.text, style=style_heading3)
            elif element.name == 'p':
                doc.add_paragraph(element.text, style=style_normal)
            elif element.name == 'pre':
                code_lines = element.text.split('\n')
                for line in code_lines:
                    p = doc.add_paragraph(line, style=style_code)
                    p.paragraph_format.left_indent = Inches(0.5)
            elif element.name == 'table':
                # Add table to document
                table_data = []
                for row in element.find_all('tr'):
                    row_data = [cell.text.strip() for cell in row.find_all(['th', 'td'])]
                    table_data.append(row_data)
                
                if table_data:
                    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                    table.style = 'Table Grid'
                    for i, row in enumerate(table_data):
                        for j, cell in enumerate(row):
                            table.cell(i, j).text = cell
                            if i == 0:  # Header row
                                run = table.cell(i, j).paragraphs[0].runs[0]
                                run.font.bold = True
                                run.font.name = 'Poppins'
                                run.font.size = Pt(10)
                                table.cell(i, j).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                # Apply header style (green background with white text)
                                shading_elm = parse_xml(r'<w:shd {} w:fill="11A64A"/>'.format(nsdecls('w')))
                                table.cell(i, j)._element.get_or_add_tcPr().append(shading_elm)
                                run.font.color.rgb = RGBColor(255, 255, 255)
                            else:
                                run = table.cell(i, j).paragraphs[0].runs[0]
                                run.font.name = 'Poppins'
                                run.font.size = Pt(9)
                                table.cell(i, j).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                # Add some space after the table
                doc.add_paragraph()
            elif element.name in ['ul', 'ol']:
                for li in element.find_all('li'):
                    paragraph = doc.add_paragraph(li.text, style=style_normal)
                    paragraph.style = 'List Bullet' if element.name == 'ul' else 'List Number'
        except Exception as e:
            print(f"Error processing element: {element}. Error: {str(e)}")
            doc.add_paragraph(str(element), style=style_normal)

    # Save to a BytesIO object
    docx_buffer = BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer

# Generate BRD button
if st.button("Generate BRD", key="generate_brd"):
    if all(st.session_state.form_fields.values()):
        progress_bar = st.progress(0)
        status_text = st.empty()

        user_types_list = st.session_state.form_fields['user_types'].split('\n')
        if 'Admin' not in user_types_list:
            user_types_list.append('Admin')
        user_types_str = '\n'.join(user_types_list)

        deliverables_list = st.session_state.form_fields['deliverables'].split('\n')
        if 'Admin Panel' not in deliverables_list:
            deliverables_list.append('Admin Panel')
        deliverables_str = '\n'.join(deliverables_list)

        part1_container = st.container()
        part2_container = st.container()
        part3_container = st.container()
        part4_container = st.container()

        # Display Lottie animation or fallback text while generating BRD
        with st.spinner("Generating your BRD..."):
            display_lottie_or_text(lottie_writing, "Generating BRD...", height=200)

        # Part 1: Confidentiality Agreement, Executive Summary, and Project Approach
        prompt_part1 = f"""
        Create the first part of a detailed Business Requirements Document (BRD) for the following project: (Numbering, heading, sub headings and paragraph should be properly formatted and don't write keyword like description or module while writing description, text sizes should be appropriate.)

        Client Name: {st.session_state.form_fields['client_name']}
        Project Description and Requirements:
        {st.session_state.form_fields['project_description']}
        Types of Users:
        {user_types_str}
        Project Deliverables:
        {deliverables_str}

        Include the following sections:
        """

        if "Confidentiality Agreement" in sections_part1:
            prompt_part1 += f"""
            1. Confidentiality Agreement: Use the following agreement:
            {confidentiality_agreement}
            """

        if "Executive Summary" in sections_part1:
            prompt_part1 += """
            2. Executive Summary: 
            - Provide a brief overview of the project, its objectives, and key stakeholders.
            - Include a detailed list of project deliverables, explaining each one concisely.
            """

        if "Project Approach" in sections_part1:
            prompt_part1 += """
            3. Project Approach: 
            - Describe the methodology, timeline, and key milestones based on the deliverables of the project.
            - Add a note mentioning that they are just for references only, final milestones will be provided in further discussions.
            """

        prompt_part1 += """
        Provide detailed and professional content for each section, incorporating all the provided information.
        Use Markdown formatting for proper structure.
        Do not include a title for the BRD itself, as it will be added separately.
        """

        with st.spinner('Generating Part 1: Confidentiality Agreement, Executive Summary, and Project Approach...'):
            with part1_container:
                model = get_model(1, mode)
                response_part1 = generate_brd_part(prompt_part1, st.empty(), model, temp_part1)
        
        progress_bar.progress(1/4)
        status_text.text("Part 1 completed. Generating Part 2...")

        # Part 2: Functional Requirements and 3rd Party Integrations
        prompt_part2 = f"""
        Create the second part of a detailed Business Requirements Document (BRD) for the following project: (Numbering, heading, sub headings and paragraph should be properly formatted and don't write keyword like description or module while writing description, text sizes should be appropriate.)

        Client Name: {st.session_state.form_fields['client_name']}
        Project Description and Requirements:
        {st.session_state.form_fields['project_description']}
        Types of Users:
        {user_types_str}
        Project Deliverables:
        {deliverables_str}

        Previously generated content:
        {response_part1}

        Now, include the following sections:
        """

        if "Functional Requirements" in sections_part2:
            prompt_part2 += """
            1. Functional Requirements:
            - For each deliverable, create at least 48 detailed modules as per requirement.
            - Structure each requirement as: Module, Sub - Module (Optional), Description.
            - Ensure the requirements are comprehensive and cover all aspects of the project.
            - Write description of each module in 3 or more lines respectively.
            - Descripton must be present against each module.
            - For each user-side module, include a corresponding management module in the admin panel.
            - Don't use any keywords like Module & Description in output.
            """

        if "3rd Party Integrations and API Suggestions" in sections_part2:
            prompt_part2 += """
            2. 3rd Party Integrations and API Suggestions:
            - Based on the functional requirements, suggest potential 3rd party integrations or APIs that could be used to implement or enhance various features of the system.
            - For each suggestion, provide:
              a. Name of the 3rd party service or API
              b. Brief description of its functionality or requirement it can address
              c. Try to suggest API or services specifically for Indian region & Mention their international alternatives too if client is from any other country.
            """

        prompt_part2 += """
        Provide detailed and professional content, incorporating all the provided information and ensuring consistency with the previously generated sections.
        Use Markdown formatting for proper structure.
        """

        with st.spinner('Generating Part 2: Functional Requirements and 3rd Party Integrations...'):
            with part2_container:
                model = get_model(2, mode)
                response_part2 = generate_brd_part(prompt_part2, st.empty(), model, temp_part2)

        progress_bar.progress(2/4)
        status_text.text("Part 2 completed. Generating Part 3...")

        # Part 3: User Journey, User Stories, and Non-Functional Requirements
        prompt_part3 = f"""
        Create the third part of a detailed Business Requirements Document (BRD) for the following project: (Numbering, heading, sub headings and paragraph should be properly formatted and don't write keyword like description or module while writing description, text sizes should be appropriate.)

        Client Name: {st.session_state.form_fields['client_name']}
        Project Description and Requirements:
        {st.session_state.form_fields['project_description']}
        Types of Users:
        {user_types_str}
        Project Deliverables:
        {deliverables_str}

        Previously generated content:
        {response_part1}
        {response_part2}

        Now, include the following sections:
        """

        if "User Journey and User Stories" in sections_part3:
            prompt_part3 += """
            1. User Journey and User Stories:
            - For each user type identified earlier, create a detailed user journey that outlines their interaction with the system from start to finish.
            - Break down each journey into individual screens or key interaction points.
            - For each screen or interaction point, provide detailed user stories in the format: "As a [user type], I want to [action] so that [benefit]."
            - Ensure that these user journeys and stories align with and reference the functional requirements generated in the previous section.
            """

        if "Non-Functional Requirements" in sections_part3:
            prompt_part3 += """
            2. Non-Functional Requirements: 
            Specify performance, security, scalability, and other non-functional aspects of the system.
            """

        prompt_part3 += """
        Provide detailed and professional content for each section, incorporating all the provided information and ensuring consistency with the previously generated sections.
        Use Markdown formatting for proper structure, including tables where specified.
        """

        with st.spinner('Generating Part 3: User Journey, User Stories, and Non-Functional Requirements...'):
            with part3_container:
                model = get_model(3, mode)
                response_part3 = generate_brd_part(prompt_part3, st.empty(), model, temp_part3)

        progress_bar.progress(3/4)
        status_text.text("Part 3 completed. Generating Part 4...")

        # Part 4: Annexure
        prompt_part4 = f"""
        Create the fourth part of a detailed Business Requirements Document (BRD) for the following project: (Numbering, heading, sub headings and paragraph should be properly formatted and don't write keyword like description or module while writing description, text sizes should be appropriate.)

        Client Name: {st.session_state.form_fields['client_name']}
        Project Description and Requirements:
        {st.session_state.form_fields['project_description']}
        Types of Users:
        {user_types_str}
        Project Deliverables:
        {deliverables_str}

        Previously generated content:
        {response_part1}
        {response_part2}
        {response_part3}

        Now, include the following section:
        """

        if "Annexure" in sections_part4:
            prompt_part4 += """
            Annexure: 
            a. Functional Requirements: Create a separate table for each deliverable in the Functional Requirements. Each table should have the following columns:
            - Requirement ID
            - Module
            - Description

            b. 3rd Party Services and APIs: Create a table summarizing all suggested 3rd party services and APIs. This table should have the following columns:
            - Service/API Name
            - Functional Area
            - Description
            """

        prompt_part4 += """
        Provide detailed and professional content for this section, incorporating all the provided information and ensuring consistency with the previously generated sections.
        Use Markdown formatting for proper structure, including tables where specified.
        """

        with st.spinner('Generating Part 4: Annexure...'):
            with part4_container:
                model = get_model(4, mode)
                response_part4 = generate_brd_part(prompt_part4, st.empty(), model, temp_part4)

        progress_bar.progress(1.0)
        status_text.text("BRD Generation Completed!")

        # Combine all parts
        full_brd = f"""
# Business Requirements Document (BRD)

**Client Name:** {st.session_state.form_fields['client_name']}

{response_part1}

{response_part2}

{response_part3}

{response_part4}
"""

        st.success("BRD Generated Successfully!")
        display_lottie_or_text(lottie_completed, "BRD Generation Completed!", height=200)

        # Download options
        col1, col2, col3 = st.columns(3)
        with col1:
            st.download_button(
                label="Download BRD as Markdown",
                data=full_brd,
                file_name=f"BRD_{st.session_state.form_fields['client_name']}.md",
                mime="text/markdown"
            )
        with col2:
            pdf_buffer = convert_markdown_to_pdf(full_brd)
            st.download_button(
                label="Download BRD as PDF",
                data=pdf_buffer,
                file_name=f"BRD_{st.session_state.form_fields['client_name']}.pdf",
                mime="application/pdf"
            )
        with col3:
            docx_buffer = convert_markdown_to_docx(full_brd)
            st.download_button(
                label="Download BRD as DOCX",
                data=docx_buffer,
                file_name=f"BRD_{st.session_state.form_fields['client_name']}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        # Display the generated BRD
        st.markdown("## Generated Business Requirements Document")
        st.markdown(full_brd)

    else:
        st.error("Please fill in all fields before generating the BRD.")

# Add some final instructions or information
st.markdown("""
---
### How to use this BRD Generator:
1. Fill in all the required fields in the form above.
2. Adjust the Advanced Settings in the sidebar if needed.
3. Click on the "Generate BRD" button to create your Business Requirements Document.
4. The generated BRD will be displayed on this page.
5. You can download the BRD as a Markdown file, PDF, or DOCX using the buttons provided.

If you need to make changes, simply update the form fields or Advanced Settings and generate the BRD again.
""")

# Footer
st.markdown("---")
st.markdown("© 2024 EMB-AI. All rights reserved.")

# Set default zoom level to 90%
st.markdown(
    """
    <script>
        document.body.style.zoom = "90%";
    </script>
    """,
    unsafe_allow_html=True
)

# Add any additional configurations or settings here
if __name__ == "__main__":
    # You can add any startup configurations or checks here
    pass
